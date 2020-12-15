from django.contrib.auth.decorators import login_required
import os
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, Table, TableStyle, Image, SimpleDocTemplate, Spacer, Flowable, BaseDocTemplate, Frame, NextPageTemplate, PageBreak, PageTemplate
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib import colors
from reportlab.rl_settings import defaultPageSize
from reportlab.lib.units import inch, cm
from reportlab import platypus
from reportlab.pdfbase.pdfmetrics import stringWidth
import csv
import docx
from docx import Document
from django.http import HttpResponse
from django.shortcuts import render

this_path = os.getcwd()
PAGE_WIDTH  = defaultPageSize[0]
PAGE_HEIGHT = defaultPageSize[1]


def piePagina(canvas,doc):
    canvas.saveState()
    canvas.setFillGray(0.3)
    canvas.setFont('Helvetica-Bold',6)
    width = doc.width
    texto1='Campus "Ingeniero Manuel Agustín Haz Álvarez", Av. Quito km. 1 1/2 vía a Santo Domingo de los Tsáchilas'
    texto2='Tel: (593) 5370-2220 - info@uteq.edu.ec - www.uteq.edu.ec'
    texto3='QUEVEDO - LOS RÍOS - ECUADOR'
    tam1 = stringWidth(texto1,'Helvetica-Bold',6)
    tam2 = stringWidth(texto2,'Helvetica-Bold',6)
    tam3 = stringWidth(texto3,'Helvetica-Bold',6)
    canvas.drawString((width-tam1) + 110, 0.7 * inch, texto1)
    canvas.drawString((width-tam2) + 110, 0.6 * inch, texto2)
    canvas.drawString((width-tam3) + 110, 0.5 * inch, texto3)
    canvas.setFillColorRGB(250 / 256, 210 / 256, 1 / 256)
    canvas.setStrokeColorRGB(250 / 256, 210 / 256, 1 / 256)
    canvas.rect(width+115, 0.5 * inch, 4, 30, fill=1)
    canvas.setFillColorRGB(76 / 256, 145 / 256, 65 / 256)
    canvas.setStrokeColorRGB(76 / 256, 145 / 256, 65 / 256)
    canvas.rect(width+122, 0.5 * inch, 8, 30, fill=1)
    canvas.restoreState()

def encabezado(canvas,doc):
    canvas.saveState()
    canvas.setFont('Times-Roman', 9)
    canvas.drawImage(this_path + '/static/images/aok/logo.png', inch, A4[1] - 90, 55, 60, mask='auto')  # imagen logo
    canvas.setStrokeColorRGB(121 / 256, 128 / 256, 129 / 256)  # choose your line color
    canvas.line(133, A4[1] - 90, 133, A4[1]-45)  # linea
    canvas.setFont('Helvetica', 12)
    canvas.setFillGray(0.3)
    canvas.drawString(140, A4[1]-64, 'Universidad')
    canvas.setFont('Helvetica-Bold', 14)
    canvas.drawString(140, A4[1]-79, 'Técnica Estatal de Quevedo')
    canvas.setFont('Helvetica-Bold', 8)
    canvas.drawString(140, A4[1]-90, 'La primera Universidad Agropecuaria del País')
    canvas.setFillColorRGB(250 / 256, 210 / 256, 1 / 256)
    canvas.setStrokeColorRGB(250 / 256, 210 / 256, 1 / 256)
    canvas.rect(50, 742, 500, 3, fill=1)
    canvas.setFillColorRGB(76 / 256, 145 / 256, 65 / 256)
    canvas.setStrokeColorRGB(76 / 256, 145 / 256, 65 / 256)
    canvas.rect(50, 742, 500, 1, fill=1)
    canvas.restoreState()

def est_reportbateria(reactivos, nombre, carrera, facultad):
    nombre2 = nombre
    nombre = nombre.replace(" ", "")
    reactivos = reactivos
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'filename='+ nombre+"-report-estudiante.pdf"
    buff = BytesIO()
    doc = BaseDocTemplate(buff, pagesize=A4)
    story = []
    #crear en el encabezado
    frame0 = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, showBoundary=0, id='normalBorde')
    frame1 = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, showBoundary=0, id='normalBorde')
    doc.addPageTemplates([ PageTemplate(id='header', frames=frame0, onPage=encabezado, onPageEnd=piePagina), PageTemplate(id='contenido', frames=frame1, onPageEnd=piePagina),])
    story.append(NextPageTemplate('contenido'))
    #HEADER2
    h1 = ParagraphStyle('texto')
    h1.textColor = 'black'
    h1.fontName = 'Helvetica-Bold'
    h1.borderPadding = 1
    h1.fontSize = 12
    h1.alignment = TA_CENTER
    facultad = Paragraph(facultad, h1)
    carrera = Paragraph(carrera + " - REACTIVOS", h1)
    story.append(Spacer(1, 30))
    story.append(facultad)
    story.append(Spacer(1,10))
    story.append(carrera)
    story.append(Spacer(1, 30))
    #HEADER3
    tabla = []
    h2 = ParagraphStyle('texto')
    h2.alignment = TA_LEFT
    h2.fontSize = 10
    h2.textColor = 'black'
    h2.fontName = 'Helvetica-Bold'
    h3 = ParagraphStyle('texto')
    h3.alignment = TA_LEFT
    h3.fontSize = 8
    h3.textColor = 'black'
    h3.fontName = 'Helvetica'
    h4=h3
    h4.fontSize = 10
    h5=ParagraphStyle('texto')
    h5.fontSize = 10
    h5.textColor = 'black'
    h5.fontName = 'Helvetica'
    h5.alignment = TA_CENTER
    tb1 = [[Paragraph('Materia/área: ',h2),Paragraph(nombre2,h4)],[Paragraph('Cantidad: ',h2),Paragraph(len(reactivos).__str__(),h4)]]
    tb1 = Table(tb1, colWidths=[1.2*inch, 6*inch])
    story.append(tb1)
    story.append(Spacer(1, 30))
    #reactivos
    index = 1
    for i in reactivos:
        tb2 = []
        tb2.append([Paragraph('<b>N°: </b>', h3), Paragraph(index.__str__(), h3)])
        tb2.append([Paragraph('<b>PREGUNTA: </b>', h3), Paragraph(i['tipo'], h3)])
        tb2.append([Paragraph('<b>NOTA: </b>', h3), Paragraph(i['nota'].__str__() + " puntos", h3)])
        # atributos
        for a in i['atributos']:
            if a['archivo']:
                tb2.append([Paragraph("<b>" + a['atributo__nombre'] + ": </b>", h3), [Paragraph(a['texto'], h3), Image(this_path+'/media/'+a['archivo'], 200,200)]])
            else:
                tb2.append([Paragraph("<b>"+a['atributo__nombre'] + ": </b>", h3), [Paragraph(a['texto'], h3)]])
        filas = len(tb2)
        if i['tipo'] !="EMPAREJAMIENTO":
            tb2.append([Paragraph('<b>OPCIONES: </b>', h3), ''])
        else:
            tab = [[[Paragraph('<b>A </b>', h5)],[Paragraph('<b>B </b>', h5)]]]
            tab = Table(tab)
            tb2.append([Paragraph('<b>OPCIONES: </b>', h3), tab])
        index2 = 1
        for a in i['opciones']:
            if i['tipo'] != "EMPAREJAMIENTO":
                texto = a['texto'].replace("br", "")
                texto  = texto.replace("<?php", "< ?php")
                if a['archivo']:
                    tb2.append([Paragraph("<b>"+index2.__str__() + ": </b>", h3), [[Paragraph(a['texto'], h3)],[Image(this_path+'/media/'+a['archivo'])]]])
                else:
                    tb2.append([Paragraph("<b>" + index2.__str__() + ": </b>", h3), [Paragraph(texto, h3)]])
            else:
                tab = [[[Paragraph(a['texto'], h3)], [Paragraph(a['texto2'], h3)]]]
                tab = Table(tab)
                tb2.append([Paragraph("<b>" + index2.__str__() + ": </b>", h3), [tab]])
            index2 += 1
        index+=1
        if i['tipo'] != "EMPAREJAMIENTO":
            styles = TableStyle([
                ('GRID', (0, 0), (-1, -1), 0.05, colors.gray),
                ('BACKGROUND', (0, 0), (0, filas - 1), colors.lavender),
                ('GRID', (0, filas), (filas, -1), 1, colors.gray),
                ('BOX', (0, 0), (1, -1), 0.1, colors.black),
                ('SPAN', (0, filas), (1, filas)),
            ])
        else:
            styles = TableStyle([
                ('GRID', (0, 0), (-1, -1), 0.05, colors.gray),
                ('BACKGROUND', (0, 0), (0, filas - 1), colors.lavender),
                ('GRID', (0, filas), (filas, -1), 1, colors.gray),
                ('BOX', (0, 0), (1, -1), 0.1, colors.black),
            ])
        tb2 = Table(tb2, colWidths=[2*inch, doc.width - inch], style=styles)
        story.append(tb2)
        story.append(Spacer(1, 30))
    doc.build(story)
    response.write(buff.getvalue())
    buff.close()
    return response

def cord_reportbateriaexamen(count, reactivos, nombre, carrera, malla, cronograma, periodo, coordinador):
    reactivos = reactivos
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'filename=' + nombre + "-report-coordinador.pdf"
    buff = BytesIO()
    doc = BaseDocTemplate(buff, pagesize=A4)
    story = []
    # crear en el encabezado
    frame0 = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, showBoundary=0, id='normalBorde')
    frame1 = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, showBoundary=0, id='normalBorde')
    doc.addPageTemplates([PageTemplate(id='header', frames=frame0, onPage=encabezado, onPageEnd=piePagina),
                          PageTemplate(id='contenido', frames=frame1, onPageEnd=piePagina), ])
    story.append(NextPageTemplate('contenido'))
    # HEADER2
    h1 = ParagraphStyle('texto')
    h1.textColor = 'black'
    h1.fontName = 'Helvetica-Bold'
    h1.borderPadding = 1
    h1.fontSize = 10
    h1.alignment = TA_LEFT
    cronograma = Paragraph("CRONOGRAMA: "+cronograma, h1)
    carrera = Paragraph("CARRERA: "+carrera, h1)
    if malla:
        malla = Paragraph("MALLA: "+malla, h1)
    periodo = Paragraph("PERIODO: "+periodo, h1)
    coordinador = Paragraph("COORDINADOR: "+coordinador, h1)
    count = Paragraph("TOTAL DE REACTIVOS: "+str(count), h1)
    story.append(Spacer(1, 30))
    story.append(cronograma)
    story.append(Spacer(1, 10))
    story.append(carrera)
    story.append(Spacer(1, 10))
    if malla:
        story.append(malla)
        story.append(Spacer(1, 10))
    story.append(periodo)
    story.append(Spacer(1, 10))
    story.append(coordinador)
    story.append(Spacer(1, 10))
    story.append(count)
    story.append(Spacer(1, 30))
    #resumen
    h3 = ParagraphStyle('texto')
    h3.alignment = TA_LEFT
    h3.fontSize = 8
    h3.textColor = 'black'
    h3.fontName = 'Helvetica'
    tb1 = []
    for i in reactivos:
        tb1.append([Paragraph(i['asignatura'],h3), Paragraph(str(i['reactivos']),h3)])
    styles = TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.05, colors.gray),
        ('GRID', (0, len(reactivos)), (len(reactivos), -1), 1, colors.gray),
        ('BOX', (0, 0), (1, -1), 0.1, colors.black),
        ('SPAN', (0, len(reactivos)), (1, len(reactivos))),
    ])
    tb2 = Table(tb1, colWidths=[270,100], style=styles)
    story.append(tb2)
    story.append(Spacer(1, 30))
    #reactivos a detalle
    doc.build(story)
    response.write(buff.getvalue())
    buff.close()
    return response

def cord_reportbateriaexamendetalle(reactivos, asignatura):
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'filename=' + asignatura + "-report-estudiante.pdf"
    buff = BytesIO()
    doc = BaseDocTemplate(buff, pagesize=A4)
    story = []
    # crear en el encabezado
    frame0 = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, showBoundary=0, id='normalBorde')
    frame1 = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, showBoundary=0, id='normalBorde')
    doc.addPageTemplates([PageTemplate(id='header', frames=frame0, onPage=encabezado, onPageEnd=piePagina), PageTemplate(id='contenido', frames=frame1, onPageEnd=piePagina), ])
    story.append(NextPageTemplate('contenido'))
    #header1
    h1 = ParagraphStyle('texto')
    h1.textColor = 'black'
    h1.fontName = 'Helvetica-Bold'
    h1.borderPadding = 1
    h1.fontSize = 12
    h1.alignment = TA_CENTER
    h3 = ParagraphStyle('texto')
    h3.alignment = TA_LEFT
    h3.fontSize = 8
    h3.textColor = 'black'
    h3.fontName = 'Helvetica'
    h5 = ParagraphStyle('texto')
    h5.fontSize = 10
    h5.textColor = 'black'
    h5.fontName = 'Helvetica'
    h5.alignment = TA_CENTER
    asignatura = Paragraph(asignatura, h1)
    story.append(Spacer(1, 30))
    story.append(asignatura)
    story.append(Spacer(1, 10))
    # reactivos
    index = 1
    for i in reactivos:
        tb2 = []
        tb2.append([Paragraph('<b>N°: </b>', h3), Paragraph(index.__str__(), h3)])
        tb2.append([Paragraph('<b>PREGUNTA: </b>', h3), Paragraph(i['tipo'], h3)])
        tb2.append([Paragraph('<b>NOTA: </b>', h3), Paragraph(i['nota'].__str__() + " puntos", h3)])
        # atributos
        for a in i['atributos']:
            if a['archivo']:
                tb2.append([Paragraph("<b>" + a['atributo__nombre'] + ": </b>", h3),
                            [Paragraph(a['texto'], h3), Image(this_path + '/media/' + a['archivo'], 200, 200)]])
            else:
                tb2.append([Paragraph("<b>" + a['atributo__nombre'] + ": </b>", h3), [Paragraph(a['texto'], h3)]])
        filas = len(tb2)
        if i['tipo'] != "EMPAREJAMIENTO":
            tb2.append([Paragraph('<b>OPCIONES: </b>', h3), ''])
        else:
            tab = [[[Paragraph('<b>A </b>', h5)], [Paragraph('<b>B </b>', h5)]]]
            tab = Table(tab)
            tb2.append([Paragraph('<b>OPCIONES: </b>', h3), tab])
        index2 = 1
        for a in i['opciones']:
            if i['tipo'] != "EMPAREJAMIENTO":
                texto = a['texto'].replace("br", "")
                texto = texto.replace("<?php", "< ?php")
                if a['archivo']:
                    tb2.append([Paragraph("<b>" + index2.__str__() + ": </b>", h3),
                                [[Paragraph(a['texto'], h3)], [Image(this_path + '/media/' + a['archivo'])]]])
                else:
                    tb2.append([Paragraph("<b>" + index2.__str__() + ": </b>", h3), [Paragraph(texto, h3)]])
            else:
                tab = [[[Paragraph(a['texto'], h3)], [Paragraph(a['texto2'], h3)]]]
                tab = Table(tab)
                tb2.append([Paragraph("<b>" + index2.__str__() + ": </b>", h3), [tab]])
            index2 += 1
        index += 1
        if i['tipo'] != "EMPAREJAMIENTO":
            styles = TableStyle([
                ('GRID', (0, 0), (-1, -1), 0.05, colors.gray),
                ('BACKGROUND', (0, 0), (0, filas - 1), colors.lavender),
                ('GRID', (0, filas), (filas, -1), 1, colors.gray),
                ('BOX', (0, 0), (1, -1), 0.1, colors.black),
                ('SPAN', (0, filas), (1, filas)),
            ])
        else:
            styles = TableStyle([
                ('GRID', (0, 0), (-1, -1), 0.05, colors.gray),
                ('BACKGROUND', (0, 0), (0, filas - 1), colors.lavender),
                ('GRID', (0, filas), (filas, -1), 1, colors.gray),
                ('BOX', (0, 0), (1, -1), 0.1, colors.black),
            ])
        tb2 = Table(tb2, colWidths=[2 * inch, doc.width - inch], style=styles)
        story.append(tb2)
        story.append(Spacer(1, 30))
    doc.build(story)
    response.write(buff.getvalue())
    buff.close()
    return response

def adm_pdfreportexamen(facultad, carrera, lista):
    nombre = carrera.replace(" ", "")
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'filename=' + nombre + "-report-examen.pdf"
    buff = BytesIO()
    doc = BaseDocTemplate(buff, pagesize=A4)
    story = []
    # crear en el encabezado
    frame0 = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, showBoundary=0, id='normalBorde')
    frame1 = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, showBoundary=0, id='normalBorde')
    doc.addPageTemplates([PageTemplate(id='header', frames=frame0, onPage=encabezado, onPageEnd=piePagina),
                          PageTemplate(id='contenido', frames=frame1, onPageEnd=piePagina), ])
    story.append(NextPageTemplate('contenido'))
    #header2
    h1 = ParagraphStyle('texto')
    h1.textColor = 'black'
    h1.fontName = 'Helvetica-Bold'
    h1.borderPadding = 1
    h1.fontSize = 12
    h1.alignment = TA_CENTER
    facultad = Paragraph(facultad, h1)
    carrera = Paragraph(carrera, h1)
    nota = Paragraph('Calificaciones de Examen Complexivo', h1)
    story.append(Spacer(1, 30))
    story.append(facultad)
    story.append(Spacer(1, 10))
    story.append(carrera)
    story.append(Spacer(1, 30))
    story.append(nota)
    story.append(Spacer(1, 5))
    #tbla de datos
    h2 = ParagraphStyle('texto')
    h2.textColor = 'black'
    h2.fontName = 'Helvetica-Bold'
    h2.borderPadding = 1
    h2.fontSize = 10
    h2.alignment = TA_LEFT
    h3 = ParagraphStyle(name='Justify', alignment=TA_JUSTIFY, fontName='Helvetica', textColor='black', fontSize=10)
    header = []
    header.append([Paragraph('<b>Cédula</b>', h2),Paragraph('<b>Apellidos y nombres</b>', h2),Paragraph('<b>Compe. Generales</b>', h2),Paragraph('<b>Compe. Específicas</b>', h2),Paragraph('<b>Calificación final</b>', h2),Paragraph('<b>Resultado</b>', h2)])
    for i in lista:
        header.append([Paragraph(i['cedula'],h3),Paragraph(i['nombre'],h3),Paragraph(i['compgeneral'],h3),Paragraph(i['compespecifica'],h3),Paragraph(i['calificacion'],h3),Paragraph(i['resultado'],h3)])
    styles = TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.black),
    ])
    tb2 = Table(header, style=styles)
    tb2._argW[1] = 2 * inch
    story.append(tb2)
    doc.build(story)
    response.write(buff.getvalue())
    buff.close()
    return response

def adm_csvreportexamen(facultad, carrera, lista):
    nombre = carrera.replace(" ", "")
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'filename=' + nombre + "-report-examen.csv"
    buff = BytesIO()
    writer = csv.writer(response)
    writer.writerow(['Cedula','Apellidos y nombres','Compe. Generales','Compe. Especificas','Califi. final','Resultado'])
    for i in lista:
        writer.writerow([i['cedula'],i['nombre'],i['compgeneral'],i['compespecifica'],i['calificacion'],i['resultado']])
    response.write(buff.getvalue())
    buff.close()
    return response

def adm_docreportexamen(facultad, carrera, lista):
    nombre = carrera.replace(" ", "")
    response = HttpResponse(content_type='application/docx')
    response['Content-Disposition'] = 'filename=' + nombre + "-report-examen.docx"
    buff = BytesIO()
    document = Document()
    document.add_heading(facultad, level=2)
    document.add_heading(carrera, level=2)
    document.add_heading('Calificaciones de Examen Complexivo', level=2)
    t = document.add_table(rows=len(lista)+1, cols=6, style='TableGrid')
    t.rows[0].cells[0].text = 'Cedula'
    t.rows[0].cells[1].text = 'Apellidos y nombres'
    t.rows[0].cells[2].text = 'Compe. Generales'
    t.rows[0].cells[3].text = 'Compe. Especificas'
    t.rows[0].cells[4].text = 'Califi. final'
    t.rows[0].cells[5].text = 'Resultado'
    index = 1
    for i in lista:
        t.rows[index].cells[0].text = i['cedula']
        t.rows[index].cells[1].text = i['nombre']
        t.rows[index].cells[2].text = i['compgeneral']
        t.rows[index].cells[3].text = i['compespecifica']
        t.rows[index].cells[4].text = i['calificacion']
        t.rows[index].cells[5].text = i['resultado']
        index+=1
    document.save(buff)
    response.write(buff.getvalue())
    buff.close()
    return response

def cord_reportrevision(tipo, facultad, carrera,lista):
    nombre = carrera.replace(" ", "")
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'filename=' + nombre + "-report-examen.pdf"
    buff = BytesIO()
    doc = BaseDocTemplate(buff, pagesize=A4)
    story = []
    # crear en el encabezado
    frame0 = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, showBoundary=0, id='normalBorde')
    frame1 = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, showBoundary=0, id='normalBorde')
    doc.addPageTemplates([PageTemplate(id='header', frames=frame0, onPage=encabezado, onPageEnd=piePagina), PageTemplate(id='contenido', frames=frame1, onPageEnd=piePagina), ])
    story.append(NextPageTemplate('contenido'))
    if tipo == "area":
        h1 = ParagraphStyle('texto')
        h1.textColor = 'black'
        h1.fontName = 'Helvetica-Bold'
        h1.borderPadding = 1
        h1.fontSize = 12
        h1.alignment = TA_CENTER
        nota = Paragraph('Asignación de reactivos generales', h1)
        story.append(Spacer(1, 30))
        story.append(nota)
        story.append(Spacer(1, 5))
    else:
        h1 = ParagraphStyle('texto')
        h1.textColor = 'black'
        h1.fontName = 'Helvetica-Bold'
        h1.borderPadding = 1
        h1.fontSize = 12
        h1.alignment = TA_CENTER
        facultad = Paragraph(facultad, h1)
        carrera = Paragraph(carrera, h1)
        nota = Paragraph('Asignación de reactivos específicos', h1)
        story.append(Spacer(1, 30))
        story.append(facultad)
        story.append(Spacer(1, 10))
        story.append(carrera)
        story.append(Spacer(1, 30))
        story.append(nota)
        story.append(Spacer(1, 5))
    #tabla
    h2 = ParagraphStyle('texto')
    h2.textColor = 'black'
    h2.fontName = 'Helvetica-Bold'
    h2.borderPadding = 1
    h2.fontSize = 10
    h2.alignment = TA_LEFT
    h3 = ParagraphStyle(name='Justify', alignment=TA_JUSTIFY, fontName='Helvetica', textColor='black', fontSize=10)
    header = []
    if tipo == "area":
        header.append([Paragraph('<b>Docente</b>', h2), Paragraph('<b>Área</b>', h2), Paragraph('<b>Cantidad de reactivos</b>', h2), Paragraph('<b>Revisión</b>', h2)])
    else:
        header.append([Paragraph('<b>Docente</b>', h2), Paragraph('<b>Asignatura</b>', h2), Paragraph('<b>Cantidad de reactivos</b>', h2), Paragraph('<b>Revisión</b>', h2)])
    for i in lista:
        if tipo == "area":
            nombre = i['asignacion'].persona.apellido1 + " " + i['asignacion'].persona.apellido2 + " " + i['asignacion'].persona.nombres
            seccion = i['asignacion'].area.nombre
        else:
            nombre = i['asignacion'].docente.persona.apellido1 + " " + i['asignacion'].docente.persona.apellido2 + " " + i['asignacion'].docente.persona.nombres
            seccion = i['asignacion'].asignatura.asignatura.nombre
        count = str(i['count']) + " de " + str(i['asignacion'].cantidad)
        if i['asignacion'].revision:
            revision = "REVISADO"
        else:
            revision = ""
        header.append([Paragraph(nombre, h3), Paragraph(seccion, h3), Paragraph(count, h3), Paragraph(revision, h3)])
    styles = TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.black),
    ])
    tb2 = Table(header, style=styles)
    tb2._argW[0] = 2 * inch
    tb2._argW[1] = 2 * inch
    story.append(tb2)
    doc.build(story)
    response.write(buff.getvalue())
    buff.close()
    return response